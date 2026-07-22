# -*- coding: utf-8 -*-
"""
公文排版工具 - 核心逻辑模块
参考 gongwen.html 重写，包含：文本规范化、文档解析、配置管理、Word 导出、文件导入
符合 GB/T 9704 标准
"""

import re
import json
import os
from copy import deepcopy
from dataclasses import dataclass, field
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Cm, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 节点类型
# ============================================================
DOCUMENT_TITLE = "DOCUMENT_TITLE"
HEADING_1 = "HEADING_1"
HEADING_2 = "HEADING_2"
HEADING_3 = "HEADING_3"
HEADING_4 = "HEADING_4"
PARAGRAPH = "PARAGRAPH"
ADDRESSEE = "ADDRESSEE"
ATTACHMENT = "ATTACHMENT"
SIGNATURE = "SIGNATURE"
DATE = "DATE"


@dataclass
class DocNode:
    """文档 AST 节点"""
    type: str
    content: str
    line_number: int = 0
    is_multiple: bool = False
    items: list = field(default_factory=list)


@dataclass
class DocAST:
    """文档抽象语法树"""
    title: Optional[DocNode] = None
    body: List[DocNode] = field(default_factory=list)


# ============================================================
# 文本规范化（对应 JS 的 l9 函数）
# ============================================================

# 半角→全角替换规则
_CJK_PATTERN = re.compile(r'([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef])\.')

_NORMALIZE_RULES = [
    (re.compile(r','), '，'),
    (_CJK_PATTERN, r'\1。'),
    (re.compile(r':'), '：'),
    (re.compile(r';'), '；'),
    (re.compile(r'\('), '（'),
    (re.compile(r'\)'), '）'),
    (re.compile(r'\?'), '？'),
    (re.compile(r'!'), '！'),
]


def normalize_text(text):
    """文本规范化：半角→全角、NBSP→空格、trim、折叠空行"""
    count = 0
    for pattern, replacement in _NORMALIZE_RULES:
        if '$1' in replacement:
            new_text, n = pattern.subn(replacement, text)
        else:
            new_text, n = pattern.subn(replacement, text)
        count += n
        text = new_text

    # NBSP → space
    text, n = re.subn('\u00A0', ' ', text)
    count += n

    # trim each line
    lines = text.split('\n')
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped != line:
            count += 1
        lines[i] = stripped
    text = '\n'.join(lines)

    # collapse 3+ newlines to 2
    text, n = re.subn(r'\n{3,}', '\n\n', text)
    count += n

    return text, count


# ============================================================
# 文档解析器（对应 JS 的 zA 函数）
# ============================================================

# 解析正则
RE_H1 = re.compile(r'^[一二三四五六七八九十]+、')
RE_H2 = re.compile(r'^[（(][一二三四五六七八九十]+[）)]')
RE_H3 = re.compile(r'^\d+[.．]')
RE_H4 = re.compile(r'^[（(]\d+[）)]')
RE_ATTACHMENT = re.compile(r'^附件[：:]')
RE_DATE = re.compile(r'^\d{4}年\d{1,2}月\d{1,2}日$')

# 结尾标点（不应出现在署名中）
END_PUNCTUATIONS = ['。', '：', ':', '；', ';', '！', '!', '？', '?', '，', ',']

# 机关关键字
ORG_KEYWORDS = [
    '人民政府', '政府', '委员会', '办公厅', '办公室',
    '党委', '党组', '部', '厅', '局', '委', '院',
    '会', '集团', '公司', '中央'
]

MAX_ADDRESSEE_LENGTH = 40


def _is_addressee(line):
    """判断是否为主送机关（以：或:结尾，长度≤40，不是标题或附件）"""
    if not (line.endswith('：') or line.endswith(':')):
        return False
    if len(line) > MAX_ADDRESSEE_LENGTH:
        return False
    if RE_H1.match(line) or RE_ATTACHMENT.match(line):
        return False
    return True


def _is_short_org_paragraph(node):
    """判断是否为短段落（可能为署名）：空或≤15字且不以标点结尾"""
    if not node or node.type != PARAGRAPH:
        return False
    content = node.content.strip()
    if len(content) == 0 or len(content) > 15:
        return False
    return not any(content.endswith(p) for p in END_PUNCTUATIONS)


def _contains_org_keyword(text):
    """检查是否包含机关关键字"""
    return any(kw in text for kw in ORG_KEYWORDS)


def _classify_line(line):
    """分类单行文本"""
    line = line.strip()
    if RE_ATTACHMENT.match(line):
        return ATTACHMENT
    if RE_DATE.match(line):
        return DATE
    if RE_H1.match(line):
        return HEADING_1
    if RE_H2.match(line):
        return HEADING_2
    if RE_H3.match(line):
        return HEADING_3
    if RE_H4.match(line):
        return HEADING_4
    return PARAGRAPH


def _parse_attachment_items(text, start_num=1):
    """解析多项目附件（如：1.xxx 2.xxx 3.xxx）"""
    items = []
    remaining = text
    num = start_num
    while remaining:
        pattern = re.compile(r'^' + str(num) + r'[.．．.]\s*')
        m = pattern.match(remaining)
        if not m:
            break
        remaining = remaining[m.end():]
        # 找下一个数字编号的位置
        next_pos = re.search(r'(?=\d+[.．．.])', remaining)
        if next_pos:
            name = remaining[:next_pos.start()].strip()
            remaining = remaining[next_pos.start():]
        else:
            name = remaining.strip()
            remaining = ''
        items.append({'index': num, 'name': name})
        num += 1
    return items, remaining


def _parse_attachment(line, lines, line_idx):
    """解析附件行（支持单项目和多项目）"""
    m = re.match(r'^附件[：:](.*)$', line)
    if not m:
        raise ValueError("Invalid attachment line")
    content = m.group(1).strip()
    # 检查是否以 1. 开头
    num_match = re.match(r'^(\d+)[.．．.]', content)
    if not num_match or num_match.group(1) != '1':
        # 单项目附件
        return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                       is_multiple=False, items=[{'index': 0, 'name': content}]), line_idx + 1

    # 多项目附件
    all_items = []
    current_text = content
    current_num = 1
    current_idx = line_idx
    while True:
        items, remaining = _parse_attachment_items(current_text, current_num)
        if remaining.strip() != '':
            # 有剩余文本，不是纯编号列表
            if current_idx == line_idx:
                return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                               is_multiple=False, items=[{'index': 0, 'name': content}]), line_idx + 1
            return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                           is_multiple=True, items=all_items), current_idx + 1
        if not items:
            if current_idx == line_idx:
                return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                               is_multiple=False, items=[{'index': 0, 'name': content}]), line_idx + 1
            return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                           is_multiple=True, items=all_items), current_idx + 1
        all_items.extend(items)
        current_num += len(items)
        next_idx = current_idx + 1
        if next_idx < len(lines):
            next_line = lines[next_idx].strip()
            if not next_line:
                current_idx = next_idx
                continue
            next_num_match = re.match(r'^(\d+)[.．．.]', next_line)
            if next_num_match and int(next_num_match.group(1)) == current_num:
                current_text = next_line
                current_idx = next_idx
                continue
        break
    return DocNode(type=ATTACHMENT, content=line, line_number=line_idx + 1,
                   is_multiple=True, items=all_items), current_idx + 1


def parse_document(text):
    """解析文档文本为 AST（对应 JS 的 zA 函数）"""
    lines = text.split('\n')
    title = None
    body = []
    has_title = False
    has_addressee = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        line_num = i + 1

        if not has_title:
            title = DocNode(type=DOCUMENT_TITLE, content=line, line_number=line_num)
            has_title = True
            i += 1
            continue

        if not has_addressee and _is_addressee(line):
            has_addressee = True
            body.append(DocNode(type=ADDRESSEE, content=line, line_number=line_num))
            i += 1
            continue

        if RE_ATTACHMENT.match(line):
            node, next_idx = _parse_attachment(line, lines, i)
            body.append(node)
            i = next_idx
            continue

        node_type = _classify_line(line)
        body.append(DocNode(type=node_type, content=line, line_number=line_num))
        i += 1

    # 后处理：短段落+含机关关键字 → 署名
    for idx in range(1, len(body)):
        if (body[idx].type == DATE and
                _is_short_org_paragraph(body[idx - 1]) and
                _contains_org_keyword(body[idx - 1].content)):
            body[idx - 1] = DocNode(type=SIGNATURE, content=body[idx - 1].content,
                                    line_number=body[idx - 1].line_number)

    # 最后一个节点如果是短段落+含机关关键字 → 署名
    if body:
        last = body[-1]
        if _is_short_org_paragraph(last) and _contains_org_keyword(last.content):
            body[-1] = DocNode(type=SIGNATURE, content=last.content,
                               line_number=last.line_number)

    return DocAST(title=title, body=body)


# ============================================================
# 默认配置（对应 JS 的 mm）
# ============================================================

def get_default_config():
    """获取默认配置（符合 GB/T 9704-2012《党政机关公文格式》国家标准）
    页边距：上3.7/下3.5/左2.8/右2.6cm
    标题：2号小标宋，行距38；正文：3号仿宋，行距28.8
    一级标题：3号黑体；二级标题：3号楷体；三四级：3号仿宋
    页码：四号 Times New Roman，— 1 — 长横线格式，奇右偶左
    版记：4号仿宋
    """
    return {
        'margins': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6,
                    'footer': 2.5},
        'title': {
            'fontFamily': '方正小标宋简体',
            'fontSize': 22,
            'lineSpacing': 38,
        },
        'body': {
            'fontFamily': '仿宋_GB2312',
            'asciiFontFamily': 'Times New Roman',
            'fontSize': 16,
            'lineSpacing': 28.8,
            'firstLineIndent': 2,
        },
        'specialOptions': {
            'boldFirstSentence': False,
            'boldHeading3': True,
            'showPageNumber': True,
            'pageNumberFont': 'Times New Roman',
            'pageNumberStyle': 'mirrored',
            'hasStamp': False,
        },
        'advanced': {
            'h1': {'fontFamily': '黑体', 'asciiFontFamily': 'Times New Roman', 'fontSize': 16},
            'h2': {'fontFamily': '楷体_GB2312', 'asciiFontFamily': 'Times New Roman', 'fontSize': 16},
            'h3': {'fontFamily': '仿宋_GB2312', 'asciiFontFamily': 'Times New Roman', 'fontSize': 16},
        },
        'header': {
            'enabled': False,
            'orgName': '',
            'docNumber': '',
            'signer': '',
        },
        'footerNote': {
            'enabled': False,
            'cc': '',
            'printer': '',
            'printDate': '',
        },
    }


def merge_config(base, patch):
    """递归合并配置"""
    result = deepcopy(base)
    for key, value in patch.items():
        if (value is not None and isinstance(value, dict) and
                isinstance(result.get(key), dict)):
            result[key] = merge_config(result[key], value)
        elif value is not None:
            result[key] = value
    return result


def normalize_config(config):
    """规范化配置，确保行距一致性"""
    title_spacing = max(config['title']['lineSpacing'], config['title']['fontSize'])
    body_spacing = max(
        config['body']['lineSpacing'],
        config['body']['fontSize'],
        config['advanced']['h1']['fontSize'],
        config['advanced']['h2']['fontSize'],
        config['advanced']['h3']['fontSize'],
    )
    if title_spacing != config['title']['lineSpacing']:
        config['title']['lineSpacing'] = title_spacing
    if body_spacing != config['body']['lineSpacing']:
        config['body']['lineSpacing'] = body_spacing
    return config


# ============================================================
# 字体选项
# ============================================================

CHINESE_FONTS = [
    ('方正小标宋_GBK', '方正小标宋_GBK'),
    ('方正小标宋简体', '方正小标宋简体'),
    ('仿宋_GB2312', '仿宋_GB2312'),
    ('仿宋', '仿宋'),
    ('黑体', '黑体'),
    ('楷体_GB2312', '楷体_GB2312'),
    ('楷体', '楷体'),
    ('宋体', '宋体'),
    ('华文中宋', '华文中宋'),
    ('Times New Roman', 'Times New Roman'),
    ('Arial', 'Arial'),
]

ASCII_FONTS = [
    ('Times New Roman', 'Times New Roman'),
    ('Arial', 'Arial'),
    ('Calibri', 'Calibri'),
    ('（跟随中文字体）', ''),
]

FONT_SIZE_NAMES = {
    42: '初号', 36: '小初', 26: '一号', 24: '小一',
    22: '二号', 18: '小二', 16: '三号', 15: '小三',
    14: '四号', 12: '小四', 10.5: '五号', 9: '小五',
}

FONT_SIZES = [42, 36, 26, 24, 22, 18, 16, 15, 14, 12, 10.5, 9]

LINE_SPACINGS = [22, 24, 26, 28, 28.8, 29, 29.6, 30, 32, 38]

INDENTS = [(0, '无缩进'), (1, '1字符'), (2, '2字符'), (3, '3字符')]

PAGE_NUMBER_STYLES = [('mirrored', '单右双左（国标）'), ('center', '全居中')]

# 字体回退表
FONT_FALLBACKS = {
    '方正小标宋简体': ['方正小标宋简体', '方正小标宋_GBK', 'FZXiaoBiaoSong-B05S', 'FZXiaoBiaoSong-B05'],
    '方正小标宋_GBK': ['方正小标宋_GBK', '方正小标宋简体', 'FZXiaoBiaoSong-B05S', 'FZXiaoBiaoSong-B05'],
    '仿宋_GB2312': ['仿宋_GB2312', '仿宋', 'FangSong_GB2312', 'FangSong', 'STFangsong'],
    '仿宋': ['仿宋', '仿宋_GB2312', 'FangSong', 'FangSong_GB2312', 'STFangsong'],
    '楷体_GB2312': ['楷体_GB2312', '楷体', 'KaiTi_GB2312', 'KaiTi', 'STKaiti'],
    '楷体': ['楷体', '楷体_GB2312', 'KaiTi', 'KaiTi_GB2312', 'STKaiti'],
    '黑体': ['黑体', 'SimHei', 'STHeiti', 'Heiti SC'],
    '宋体': ['宋体', 'SimSun', 'STSong', 'Songti SC'],
    '新宋体': ['新宋体', 'NSimSun', '宋体', 'SimSun', 'STSong'],
    '华文中宋': ['华文中宋', 'STZhongsong'],
    '华文仿宋': ['华文仿宋', 'STFangsong', '仿宋', 'FangSong'],
    '华文楷体': ['华文楷体', 'STKaiti', '楷体', 'KaiTi'],
    '华文彩云': ['华文彩云', 'STCaiyun'],
}


def get_font_fallbacks(font_name):
    """获取字体回退列表"""
    name = font_name.strip() if font_name else ''
    if not name:
        return ['Times New Roman']
    if name in FONT_FALLBACKS:
        return list(FONT_FALLBACKS[name])
    return [name]


def format_font_name_for_docx(font_name):
    """格式化字体名称用于 docx"""
    fallbacks = get_font_fallbacks(font_name)
    return ', '.join(f'"{f}"' for f in fallbacks if f)


# ============================================================
# 配置持久化
# ============================================================

CONFIG_KEY = 'docx-document-config'
TEXT_KEY = 'docx-editor-text'

# 配置文件路径
_APP_DATA_DIR = os.path.join(os.environ.get('APPDATA', os.path.expanduser('~')), 'GongwenTool')
_CONFIG_FILE = os.path.join(_APP_DATA_DIR, 'config.json')
_TEXT_FILE = os.path.join(_APP_DATA_DIR, 'editor_text.txt')


def save_config(config):
    """保存配置到文件"""
    try:
        os.makedirs(_APP_DATA_DIR, exist_ok=True)
        with open(_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_config():
    """从文件加载配置"""
    try:
        if os.path.exists(_CONFIG_FILE):
            with open(_CONFIG_FILE, 'r', encoding='utf-8') as f:
                saved = json.load(f)
            merged = merge_config(get_default_config(), saved)
            return normalize_config(merged)
    except Exception:
        pass
    return normalize_config(get_default_config())


def save_text(text):
    """保存编辑器文本"""
    try:
        os.makedirs(_APP_DATA_DIR, exist_ok=True)
        with open(_TEXT_FILE, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception:
        pass


def load_text():
    """加载编辑器文本"""
    try:
        if os.path.exists(_TEXT_FILE):
            with open(_TEXT_FILE, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception:
        pass
    return ''


# ============================================================
# Word 导出（对应 JS 的 JC 函数）
# ============================================================

def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_paragraph_spacing(paragraph, line_spacing, before=0, after=0):
    """设置段落行距（精确值），段前段后默认0"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(line_spacing * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), str(int(before)))
    spacing.set(qn('w:after'), str(int(after)))
    # 明确设置段前段后为0字符（避免继承样式）
    spacing.set(qn('w:beforeLines'), '0')
    spacing.set(qn('w:afterLines'), '0')


def _set_zero_indents(paragraph):
    """设置段落左右缩进为0字符"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:right'), '0')
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:rightChars'), '0')


def _set_paragraph_border(paragraph, position='bottom', size='15', color='E00000'):
    """设置段落边框"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)
    pBdr.append(border)
    pPr.append(pBdr)


# 标点符号正则（中文和英文标点）
_PUNCTUATION_PATTERN = re.compile(
    r'[\u3000-\u303f\uff00-\uffef'
    r'，。；：、！？""''（）【】《》〈〉「」『』·…—'
    r',.!?;:()\[\]{}<>'
    r'．．]'
)


def _add_run_with_font(paragraph, text, font_ea, font_ascii, size_pt, bold=False, color=None):
    """添加带有正确字体设置的 run
    标点符号统一使用仿宋_GB2312 字体
    """
    # 将文本按标点符号拆分
    segments = []
    last_end = 0
    for m in _PUNCTUATION_PATTERN.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False))
        segments.append((m.group(), True))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False))
    if not segments:
        segments = [(text, False)]

    runs = []
    for seg_text, is_punct in segments:
        if not seg_text:
            continue
        run = paragraph.add_run(seg_text)
        run.font.size = Pt(size_pt)
        run.bold = bold
        if color:
            run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).RGBColor.from_string(color)

        # 设置字体
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        if is_punct:
            # 标点符号统一使用仿宋_GB2312
            rFonts.set(qn('w:ascii'), '仿宋_GB2312')
            rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
            rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            rFonts.set(qn('w:cs'), '仿宋_GB2312')
        else:
            rFonts.set(qn('w:ascii'), font_ascii)
            rFonts.set(qn('w:hAnsi'), font_ascii)
            rFonts.set(qn('w:eastAsia'), font_ea)
            rFonts.set(qn('w:cs'), font_ascii)
        runs.append(run)
    return runs


def _add_empty_paragraph(doc, config, count=1):
    """添加空行"""
    for _ in range(count):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, config['body']['lineSpacing'])


def _add_title(doc, title_node, config):
    """添加公文标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, config['title']['lineSpacing'])
    _set_zero_indents(p)
    font = config['title']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    _add_run_with_font(p, title_node.content, font, ascii_font,
                       config['title']['fontSize'], bold=True)
    # 标题与正文之间空一行
    _add_empty_paragraph(doc, config, count=1)


def _add_body_paragraph(doc, node, config, is_first):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    line_spacing = config['body']['lineSpacing']
    indent = config['body']['firstLineIndent']
    _set_paragraph_spacing(p, line_spacing)
    _set_zero_indents(p)

    if indent > 0:
        p.paragraph_format.first_line_indent = Pt(indent * config['body']['fontSize'])

    body_font = config['body']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    body_size = config['body']['fontSize']

    if node.type == HEADING_1:
        font = config['advanced']['h1']['fontFamily']
        ascii_font = config['advanced']['h1']['asciiFontFamily']
        size = config['advanced']['h1']['fontSize']
    elif node.type == HEADING_2:
        font = config['advanced']['h2']['fontFamily']
        ascii_font = config['advanced']['h2']['asciiFontFamily']
        size = config['advanced']['h2']['fontSize']
    elif node.type == HEADING_3:
        font = config['advanced']['h3']['fontFamily']
        ascii_font = config['advanced']['h3']['asciiFontFamily']
        size = config['advanced']['h3']['fontSize']
    elif node.type == HEADING_4:
        font = body_font
        size = body_size
    else:
        font = body_font
        size = body_size

    bold = False
    if node.type == HEADING_3 and config['specialOptions']['boldHeading3']:
        bold = True

    _add_run_with_font(p, node.content, font, ascii_font, size, bold=bold)


def _add_addressee(doc, node, config, after_title):
    """添加主送机关"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p, config['body']['lineSpacing'])
    _set_zero_indents(p)
    _add_run_with_font(p, node.content,
                       config['body']['fontFamily'],
                       config['body']['asciiFontFamily'],
                       config['body']['fontSize'])


def _add_attachment(doc, node, config, is_first):
    """添加附件"""
    indent = config['body']['firstLineIndent']
    body_font = config['body']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    body_size = config['body']['fontSize']

    if not node.is_multiple:
        # 单项目附件
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, config['body']['lineSpacing'])
        _set_zero_indents(p)
        if indent > 0:
            p.paragraph_format.left_indent = Pt(indent * body_size + 3 * body_size)
            p.paragraph_format.first_line_indent = Pt(-3 * body_size)
        _add_run_with_font(p, node.content, body_font, ascii_font, body_size)
    else:
        # 多项目附件
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, config['body']['lineSpacing'])
        _set_zero_indents(p)
        prefix = '附件：'
        _add_run_with_font(p, prefix, body_font, ascii_font, body_size)
        if indent > 0:
            p.paragraph_format.left_indent = Pt(indent * body_size + 3 * body_size)
            p.paragraph_format.first_line_indent = Pt(-3 * body_size)

        for i, item in enumerate(node.items):
            text = f"{item['index']}. {item['name']}"
            _add_run_with_font(p, text, body_font, ascii_font, body_size)


def _add_signature_and_date(doc, sig_node, date_node, config):
    """添加署名和日期（特殊布局）"""
    line_spacing = config['body']['lineSpacing']
    body_font = config['body']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    body_size = config['body']['fontSize']

    # 添加空行
    _add_empty_paragraph(doc, config, count=2)

    # 署名（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p, line_spacing)
    _set_zero_indents(p)
    _add_run_with_font(p, sig_node.content, body_font, ascii_font, body_size)

    # 日期（右对齐，右空四字）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p, line_spacing)
    _set_zero_indents(p)
    p.paragraph_format.right_indent = Pt(4 * body_size)
    _add_run_with_font(p, date_node.content, body_font, ascii_font, body_size)


def _add_header(doc, config):
    """添加版头"""
    if not (config['header']['enabled'] and config['header']['orgName']):
        return

    body_font = config['body']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    body_size = config['body']['fontSize']
    line_spacing = config['body']['lineSpacing']
    signer = config['header']['signer']
    doc_number = config['header']['docNumber']

    # 发文机关标志（红色大字）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, 48)
    _add_run_with_font(p, config['header']['orgName'],
                       '方正小标宋_GBK', 'Times New Roman', 30, color='E00000')

    # 两个空行
    _add_empty_paragraph(doc, config, count=2)

    if signer:
        # 有签发人：表格布局（左：发文字号，右：签发人）
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(8)
        table.columns[1].width = Cm(8)

        # 左侧：发文字号
        cell = table.cell(0, 0)
        _set_cell_border(cell, top='nil', left='nil', bottom='nil', right='nil')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, line_spacing)
        if doc_number:
            _add_run_with_font(p, doc_number, body_font, ascii_font, body_size)

        # 右侧：签发人
        cell = table.cell(0, 1)
        _set_cell_border(cell, top='nil', left='nil', bottom='nil', right='nil')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(p, line_spacing)
        _add_run_with_font(p, '签发人：', body_font, ascii_font, body_size)
        _add_run_with_font(p, signer, '楷体_GB2312', ascii_font, body_size)
    elif doc_number:
        # 无签发人，只有发文字号（居中）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, line_spacing)
        _add_run_with_font(p, doc_number, body_font, ascii_font, body_size)

    # 红色分隔线
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line_spacing)
    _set_paragraph_border(p, position='bottom', size='15', color='E00000')


def _add_footer_note(doc, config):
    """添加版记"""
    if not config['footerNote']['enabled']:
        return

    body_font = config['body']['fontFamily']
    ascii_font = config['body']['asciiFontFamily']
    line_spacing = config['body']['lineSpacing']

    # 上分隔线
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line_spacing)
    _set_paragraph_border(p, position='bottom', size='6', color='000000')

    # 抄送
    if config['footerNote']['cc']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, line_spacing)
        _add_run_with_font(p, '抄送：', body_font, ascii_font, 14)
        _add_run_with_font(p, config['footerNote']['cc'], body_font, ascii_font, 14)

    # 中分隔线
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line_spacing)
    _set_paragraph_border(p, position='bottom', size='3', color='000000')

    # 印发机关和日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    cell = table.cell(0, 0)
    _set_cell_border(cell, top='nil', left='nil', bottom='nil', right='nil')
    p = cell.paragraphs[0]
    _set_paragraph_spacing(p, line_spacing)
    _add_run_with_font(p, config['footerNote']['printer'] or '', body_font, ascii_font, 14)

    cell = table.cell(0, 1)
    _set_cell_border(cell, top='nil', left='nil', bottom='nil', right='nil')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p, line_spacing)
    _add_run_with_font(p, config['footerNote']['printDate'] or '', body_font, ascii_font, 14)

    # 下分隔线
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line_spacing)
    _set_paragraph_border(p, position='bottom', size='6', color='000000')


def _add_page_number(doc, config):
    """添加页码（GB/T 9704-2012：四号宋体，－１－格式，奇右偶左）"""
    if not config['specialOptions']['showPageNumber']:
        return

    style = config['specialOptions']['pageNumberStyle']
    font = config['specialOptions']['pageNumberFont']

    for section in doc.sections:
        # 国标要求奇偶页不同（外侧对齐）
        section.different_first_page_header_footer = False
        # 启用奇偶页不同的页脚
        sectPr = section._sectPr
        titlePg = sectPr.find(qn('w:titlePg'))
        if titlePg is None:
            titlePg = OxmlElement('w:titlePg')
            sectPr.append(titlePg)

        footer = section.footer
        footer.is_linked_to_previous = False

        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]

        if style == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 单右双左：奇数页右，偶数页左
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 页码格式：— 1 —（前后加长横线）
        # 前长横线
        run_dash_before = p.add_run('— ')
        run_dash_before.font.size = Pt(14)
        _set_run_font(run_dash_before, font)

        # 页码字段
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

        # 后长横线
        run_dash_after = p.add_run(' —')
        run_dash_after.font.size = Pt(14)
        _set_run_font(run_dash_after, font)

        # 设置字体（四号宋体）
        for r in [run, run2, run3]:
            r.font.size = Pt(14)
            _set_run_font(r, font)


def _set_run_font(run, font):
    """设置 run 的字体"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), font)


def _set_page_margins(doc, config):
    """设置页面边距（GB/T 9704-2012：上3.7/下3.5/左2.8/右2.6，页脚2.5cm）"""
    for section in doc.sections:
        section.top_margin = Cm(config['margins']['top'])
        section.bottom_margin = Cm(config['margins']['bottom'])
        section.left_margin = Cm(config['margins']['left'])
        section.right_margin = Cm(config['margins']['right'])
        if 'footer' in config['margins']:
            section.footer_distance = Cm(config['margins']['footer'])


def export_to_docx(ast, config, file_path, progress_callback=None):
    """导出为 Word 文档"""
    if progress_callback:
        progress_callback('正在生成文档结构…')

    doc = Document()

    # 设置页面边距
    _set_page_margins(doc, config)

    # 添加版头
    _add_header(doc, config)

    # 添加标题
    if ast.title:
        _add_title(doc, ast.title, config)

    # 遍历正文节点
    for i, node in enumerate(ast.body):
        is_first = (i == 0)

        if node.type == SIGNATURE:
            # 检查是否后跟日期
            if i + 1 < len(ast.body) and ast.body[i + 1].type == DATE:
                _add_signature_and_date(doc, node, ast.body[i + 1], config)
                # 跳过下一个日期节点
                continue
            else:
                # 普通署名
                _add_empty_paragraph(doc, config, count=2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_paragraph_spacing(p, config['body']['lineSpacing'])
                _add_run_with_font(p, node.content,
                                   config['body']['fontFamily'],
                                   config['body']['asciiFontFamily'],
                                   config['body']['fontSize'])
            continue

        if node.type == DATE:
            # 独立日期（不跟在署名后）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_spacing(p, config['body']['lineSpacing'])
            _set_zero_indents(p)
            p.paragraph_format.right_indent = Pt(4 * config['body']['fontSize'])
            _add_run_with_font(p, node.content,
                               config['body']['fontFamily'],
                               config['body']['asciiFontFamily'],
                               config['body']['fontSize'])
            continue

        if node.type == ATTACHMENT:
            _add_attachment(doc, node, config, is_first)
            continue

        if node.type == ADDRESSEE:
            _add_addressee(doc, node, config, after_title=(i == 0))
            continue

        # 普通段落和标题
        _add_body_paragraph(doc, node, config, is_first)

    # 添加版记
    _add_footer_note(doc, config)

    # 添加页码
    _add_page_number(doc, config)

    if progress_callback:
        progress_callback('正在保存文件…')

    doc.save(file_path)


# ============================================================
# 文件导入
# ============================================================

def import_file(file_path):
    """导入 .docx 或 .txt 文件，返回文本"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return _import_docx(file_path)
    elif ext == '.txt':
        return _import_txt(file_path)
    elif ext in ('.doc', '.wps'):
        raise ValueError('不支持 .doc/.wps 格式，请先用 WPS 或 Word 另存为 .docx 文件')
    else:
        raise ValueError('不支持的文件格式，仅支持 .docx 和 .txt 文件')


def _import_docx(file_path):
    """从 .docx 文件提取文本"""
    import mammoth
    with open(file_path, 'rb') as f:
        result = mammoth.extract_raw_text(f)
    text = result.value
    return _clean_text(text)


def _import_txt(file_path):
    """从 .txt 文件读取文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return _clean_text(text)


def _clean_text(text):
    """清理导入的文本"""
    lines = text.split('\n')
    cleaned = [line.strip() for line in lines]
    result = '\n'.join(cleaned)
    result = re.sub(r'\n{2,}', '\n\n', result).strip()
    return result
